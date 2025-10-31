from __future__ import annotations

import os
import re
from typing import Dict, List, Optional, Any

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    ListFlowable,
    ListItem,
    Table,
    TableStyle,
    PageBreak,
)
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus.tableofcontents import TableOfContents


def _build_filename(output_dir: str, slug: str, timestamp: str, ext: str) -> str:
    return os.path.join(output_dir, f"{slug}_{timestamp}.{ext}")


# -------- Common helpers --------

def _markdown_inline_to_rml(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"<i>\1</i>", text)
    return text


# -------- TXT Export --------

def export_txt(
    output_dir: str,
    slug: str,
    timestamp: str,
    metadata: Dict[str, object],
    heading: str,
    sections: List[Dict[str, str]],
) -> str:
    path = _build_filename(output_dir, slug, timestamp, "txt")

    lines: List[str] = []
    lines.append("SEO Metadata")
    lines.append(f"Title Tag: {metadata['title_tag']}")
    lines.append(f"Meta Description: {metadata['meta_description']}")
    lines.append("Primary Keywords: " + ", ".join(metadata["primary_keywords"]))
    lines.append("Secondary Keywords: " + ", ".join(metadata["secondary_keywords"]))
    lines.append(f"URL Slug: {metadata['url_slug']}")
    lines.append("")
    lines.append(f"H1: {heading}")
    lines.append("")

    for s in sections:
        prefix = "H2:" if s["level"] == "h2" else "H3:"
        lines.append(f"{prefix} {s['title']}")
        lines.append(s["text"].strip())
        lines.append("")

    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines).strip() + "\n")

    return path


# -------- DOCX Export --------

def _docx_add_toc(document: Document) -> None:
    p = document.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)


def export_docx(
    output_dir: str,
    slug: str,
    timestamp: str,
    metadata: Dict[str, object],
    heading: str,
    sections: List[Dict[str, str]],
    style_options: Optional[Dict[str, Any]] = None,
) -> str:
    path = _build_filename(output_dir, slug, timestamp, "docx")

    style_options = style_options or {}
    base_font_size: int = int(style_options.get("base_font_size", 11))
    h1_size: int = int(style_options.get("h1_size", 18))
    h2_size: int = int(style_options.get("h2_size", 14))
    h3_size: int = int(style_options.get("h3_size", 12))
    font_family: str = str(style_options.get("font_family", "Open Sans"))
    polish: bool = bool(style_options.get("polish", False))

    document = Document()

    # Margins: 1 inch
    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # TOC first
    document.add_heading('Table of Contents', level=1)
    _docx_add_toc(document)
    document.add_page_break()

    # Title Style (centered, bold)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(heading)
    run.bold = True
    run.font.name = font_family
    run.font.size = Pt(h1_size if not polish else 20)

    def _format_para(p: Any):
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in p.runs:
            r.font.name = font_family
            r.font.size = Pt(base_font_size if not polish else 11)

    def _add_heading(text: str):
        p = document.add_paragraph()
        r = p.add_run(text)
        r.font.name = font_family
        r.bold = True
        r.font.size = Pt(h2_size if not polish else 15)
        return p

    def _add_subheading(text: str):
        p = document.add_paragraph()
        r = p.add_run(text)
        r.font.name = font_family
        r.bold = True
        if polish:
            r.italic = True  # semi-bold approximation
            r.font.size = Pt(13)
        else:
            r.font.size = Pt(h3_size)
        return p

    # Metadata
    for line in [
        f"Title Tag: {metadata['title_tag']}",
        f"Meta Description: {metadata['meta_description']}",
        "Primary Keywords: " + ", ".join(metadata["primary_keywords"]),
        "Secondary Keywords: " + ", ".join(metadata["secondary_keywords"]),
        f"URL Slug: {metadata['url_slug']}",
    ]:
        p = document.add_paragraph(line)
        _format_para(p)

    # Content
    for s in sections:
        if s["level"] == "h2":
            _add_heading(s["title"])  # Heading Style
        else:
            _add_subheading(s["title"])  # Subheading Style
        # paragraphs and bullets
        for para in s["text"].split("\n"):
            stripped = para.strip()
            if not stripped:
                continue
            if stripped.startswith('- '):
                # bullet
                p = document.add_paragraph(stripped[2:], style=None)
                p.style = document.styles['List Bullet'] if 'List Bullet' in document.styles else None
                _format_para(p)
            else:
                p = document.add_paragraph(stripped)
                _format_para(p)

    document.save(path)
    return path


# -------- PDF Export --------

def _register_open_sans_if_available() -> Optional[str]:
    here = os.path.dirname(os.path.dirname(__file__))
    fonts_dir = os.path.join(here, 'fonts')
    regular = os.path.join(fonts_dir, 'OpenSans-Regular.ttf')
    bold = os.path.join(fonts_dir, 'OpenSans-Bold.ttf')
    italic = os.path.join(fonts_dir, 'OpenSans-Italic.ttf')
    bolditalic = os.path.join(fonts_dir, 'OpenSans-BoldItalic.ttf')
    if all(os.path.exists(p) for p in [regular, bold, italic, bolditalic]):
        pdfmetrics.registerFont(TTFont('OpenSans', regular))
        pdfmetrics.registerFont(TTFont('OpenSans-Bold', bold))
        pdfmetrics.registerFont(TTFont('OpenSans-Italic', italic))
        pdfmetrics.registerFont(TTFont('OpenSans-BoldItalic', bolditalic))
        return 'OpenSans'
    return None


def _mk_styles(font_family: str, base_font_size: int, h1_size: int, h2_size: int, h3_size: int):
    styles = getSampleStyleSheet()

    styles["BodyText"].fontName = font_family
    styles["BodyText"].fontSize = base_font_size
    styles["BodyText"].leading = int(base_font_size * 1.15)
    styles["BodyText"].spaceAfter = 6
    styles["BodyText"].alignment = TA_JUSTIFY

    styles["Title"].fontName = font_family + ("-Bold" if 'OpenSans' in font_family else "")
    styles["Title"].fontSize = h1_size

    styles["Heading1"].fontName = font_family + ("-Bold" if 'OpenSans' in font_family else "")
    styles["Heading1"].fontSize = h1_size
    styles["Heading1"].spaceAfter = 6

    styles["Heading2"].fontName = font_family + ("-Bold" if 'OpenSans' in font_family else "")
    styles["Heading2"].fontSize = h2_size
    styles["Heading2"].spaceAfter = 6

    styles["Heading3"].fontName = font_family + ("-BoldItalic" if 'OpenSans' in font_family else "")
    styles["Heading3"].fontSize = h3_size
    styles["Heading3"].spaceAfter = 6

    quote_style = ParagraphStyle(
        name="BlockQuote",
        parent=styles["BodyText"],
        leftIndent=18,
        textColor=colors.grey,
    )
    styles.add(quote_style)

    bullet_style = ParagraphStyle(
        name="BulletItem",
        parent=styles["BodyText"],
        leftIndent=12,
    )
    styles.add(bullet_style)

    return styles


def _render_markdown_like(text: str, styles) -> List[Any]:
    lines = text.splitlines()
    story: List[Any] = []

    def flush_list(bullets: List[str]):
        if not bullets:
            return []
        flow = [ListItem(Paragraph(_markdown_inline_to_rml(b), styles['BulletItem'])) for b in bullets]
        return [ListFlowable(flow, bulletType='bullet'), Spacer(1, 6)]

    def flush_table(rows: List[List[str]]):
        if not rows:
            return []
        tbl = Table(rows, hAlign='LEFT')
        tbl.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('BACKGROUND', (0,0), (-1,0), colors.whitesmoke),
            ('LEFTPADDING', (0,0), (-1,-1), 6),
            ('RIGHTPADDING', (0,0), (-1,-1), 6),
            ('TOPPADDING', (0,0), (-1,-1), 4),
            ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ]))
        return [tbl, Spacer(1, 8)]

    bullets: List[str] = []
    current_table: List[List[str]] = []

    i = 0
    while i < len(lines):
        stripped = lines[i].strip()

        if '|' in stripped and stripped.startswith('|') and stripped.endswith('|'):
            if bullets:
                story.extend(flush_list(bullets)); bullets = []
            current_table.append([c.strip() for c in stripped.strip('|').split('|')])
            i += 1
            continue
        if current_table and not (stripped.startswith('|') and stripped.endswith('|')):
            story.extend(flush_table(current_table)); current_table = []

        if stripped.startswith('### '):
            if bullets:
                story.extend(flush_list(bullets)); bullets = []
            story.append(Paragraph(_markdown_inline_to_rml(stripped[4:]), styles['Heading3']))
            story.append(Spacer(1, 6))
            i += 1
            continue
        if stripped.startswith('## '):
            if bullets:
                story.extend(flush_list(bullets)); bullets = []
            story.append(Paragraph(_markdown_inline_to_rml(stripped[3:]), styles['Heading2']))
            story.append(Spacer(1, 8))
            i += 1
            continue
        if stripped.startswith('# '):
            if bullets:
                story.extend(flush_list(bullets)); bullets = []
            story.append(Paragraph(_markdown_inline_to_rml(stripped[2:]), styles['Heading1']))
            story.append(Spacer(1, 10))
            i += 1
            continue

        if stripped.startswith('> '):
            if bullets:
                story.extend(flush_list(bullets)); bullets = []
            story.append(Paragraph(_markdown_inline_to_rml(stripped[2:]), styles['BlockQuote']))
            story.append(Spacer(1, 6))
            i += 1
            continue

        if stripped.startswith('- ') or stripped.startswith('* '):
            bullets.append(stripped[2:])
            i += 1
            continue
        else:
            if bullets:
                story.extend(flush_list(bullets)); bullets = []

        if stripped:
            story.append(Paragraph(_markdown_inline_to_rml(stripped), styles['BodyText']))
            story.append(Spacer(1, 6))
        else:
            story.append(Spacer(1, 6))
        i += 1

    if current_table:
        story.extend(flush_table(current_table))
    if bullets:
        story.extend(flush_list(bullets))

    return story


def export_pdf(
    output_dir: str,
    slug: str,
    timestamp: str,
    metadata: Dict[str, object],
    heading: str,
    sections: List[Dict[str, str]],
    style_options: Optional[Dict[str, Any]] = None,
) -> str:
    path = _build_filename(output_dir, slug, timestamp, "pdf")

    style_options = style_options or {}
    base_font_size: int = int(style_options.get("base_font_size", 11))
    h1_size: int = int(style_options.get("h1_size", 18))
    h2_size: int = int(style_options.get("h2_size", 14))
    h3_size: int = int(style_options.get("h3_size", 12))
    polish: bool = bool(style_options.get("polish", False))

    # Try Open Sans; for polish mode, use Times-Roman
    if polish:
        family = 'Times-Roman'
    else:
        family = _register_open_sans_if_available() or 'Helvetica'

    doc = SimpleDocTemplate(
        path,
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )
    styles = _mk_styles(family, base_font_size if not polish else 11, h1_size if not polish else 20, h2_size if not polish else 15, h3_size if not polish else 13)

    story: List[Any] = []

    # TOC
    story.append(Paragraph('Table of Contents', styles['Heading1']))
    toc = TableOfContents()
    toc.levelStyles = [
        ParagraphStyle(fontName=family, fontSize=12, name='TOCHeading1', leftIndent=12, firstLineIndent=-12),
        ParagraphStyle(fontName=family, fontSize=11, name='TOCHeading2', leftIndent=24, firstLineIndent=-12),
        ParagraphStyle(fontName=family, fontSize=10, name='TOCHeading3', leftIndent=36, firstLineIndent=-12),
    ]
    story.append(toc)
    story.append(PageBreak())

    # Title (centered)
    title_style = ParagraphStyle(name='CenteredTitle', parent=styles['Title'])
    title_style.alignment = 1  # center
    story.append(Paragraph(_markdown_inline_to_rml(heading), title_style))
    story.append(Spacer(1, 12))

    # Metadata
    for line in [
        f"Title Tag: {metadata['title_tag']}",
        f"Meta Description: {metadata['meta_description']}",
        "Primary Keywords: " + ", ".join(metadata["primary_keywords"]),
        "Secondary Keywords: " + ", ".join(metadata["secondary_keywords"]),
        f"URL Slug: {metadata['url_slug']}",
    ]:
        story.append(Paragraph(_markdown_inline_to_rml(line), styles['BodyText']))
    story.append(Spacer(1, 12))

    # Content
    for s in sections:
        if s["level"] == "h2":
            h = Paragraph(_markdown_inline_to_rml(s["title"]), styles['Heading2'])
            h._bookmarkName = s["title"]
            story.append(h)
        else:
            h = Paragraph(_markdown_inline_to_rml(s["title"]), styles['Heading3'])
            h._bookmarkName = s["title"]
            story.append(h)

        if polish:
            # In polish mode, treat lines as body or bullets; no markdown headings parsing
            for line in s["text"].split("\n"):
                t = line.strip()
                if not t:
                    continue
                if t.startswith('- '):
                    story.append(ListFlowable([ListItem(Paragraph(_markdown_inline_to_rml(t[2:]), styles['BulletItem']))], bulletType='bullet'))
                else:
                    story.append(Paragraph(_markdown_inline_to_rml(t), styles['BodyText']))
            story.append(Spacer(1, 12))
        else:
            story.extend(_render_markdown_like(s["text"], styles))
            story.append(Spacer(1, 12))

    def _add_page_number(canvas_obj, doc_obj):
        page_num_text = f"{canvas_obj.getPageNumber()}"
        canvas_obj.setFont('Helvetica', 9)
        canvas_obj.drawCentredString(4.25 * inch, 0.5 * inch, page_num_text)

    doc.build(story, onFirstPage=_add_page_number, onLaterPages=_add_page_number)

    return path
